"""
Enterprise AI Assistant - Document Loader
Endüstri Standartlarında Kurumsal AI Çözümü

Çoklu format döküman yükleme - PDF, DOCX, XLSX, TXT, MD, HTML desteği.
"""

from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import hashlib


class Document:
    """Döküman veri yapısı."""
    
    def __init__(
        self,
        content: str,
        metadata: Optional[Dict[str, Any]] = None,
    ):
        self.content = content
        self.metadata = metadata or {}
        self.id = self._generate_id()
    
    def _generate_id(self) -> str:
        """Unique döküman ID'si oluştur."""
        content_hash = hashlib.md5(self.content.encode()).hexdigest()[:12]
        return f"doc_{content_hash}"
    
    def __repr__(self) -> str:
        return f"Document(id={self.id}, length={len(self.content)})"


class DocumentLoader:
    """
    Döküman yükleme sınıfı - Endüstri standartlarına uygun.
    
    Desteklenen formatlar:
    - PDF (.pdf)
    - Word (.docx)
    - Excel (.xlsx, .csv)
    - Text (.txt, .md)
    - HTML (.html, .htm)
    - JSON (.json)
    """
    
    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
        ".xlsx": "xlsx",
        ".xls": "xls",
        ".csv": "csv",
        ".txt": "text",
        ".md": "markdown",
        ".html": "html",
        ".htm": "html",
        ".json": "json",
        ".xml": "xml",
    }
    
    def __init__(self):
        self._loaders = {}
    
    def load_file(self, file_path: str) -> List[Document]:
        """
        Dosyayı yükle ve Document listesi döndür.
        
        Args:
            file_path: Dosya yolu
            
        Returns:
            Document listesi
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Dosya bulunamadı: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Desteklenmeyen format: {extension}")
        
        file_type = self.SUPPORTED_EXTENSIONS[extension]
        
        # Load based on file type
        if file_type == "pdf":
            return self._load_pdf(path)
        elif file_type == "docx":
            return self._load_docx(path)
        elif file_type in ("xlsx", "xls"):
            return self._load_excel(path)
        elif file_type == "csv":
            return self._load_csv(path)
        elif file_type in ("text", "markdown"):
            return self._load_text(path)
        elif file_type == "html":
            return self._load_html(path)
        elif file_type == "json":
            return self._load_json(path)
        else:
            return self._load_text(path)
    
    def load_directory(
        self,
        directory_path: str,
        recursive: bool = True,
        extensions: Optional[List[str]] = None,
    ) -> List[Document]:
        """
        Klasördeki tüm dökümanları yükle.
        
        Args:
            directory_path: Klasör yolu
            recursive: Alt klasörleri de tara
            extensions: Sadece bu uzantıları yükle
            
        Returns:
            Document listesi
        """
        path = Path(directory_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Klasör bulunamadı: {directory_path}")
        
        documents = []
        
        # Filter extensions
        valid_extensions = extensions or list(self.SUPPORTED_EXTENSIONS.keys())
        
        # Get files
        pattern = "**/*" if recursive else "*"
        files = list(path.glob(pattern))
        
        for file_path in files:
            if file_path.is_file() and file_path.suffix.lower() in valid_extensions:
                try:
                    docs = self.load_file(str(file_path))
                    documents.extend(docs)
                    print(f"✅ Yüklendi: {file_path.name}")
                except Exception as e:
                    print(f"⚠️ Yüklenemedi: {file_path.name} - {e}")
        
        return documents
    
    def _get_base_metadata(self, path: Path) -> Dict[str, Any]:
        """Temel metadata oluştur."""
        stat = path.stat()
        return {
            "source": str(path),
            "filename": path.name,
            "file_type": path.suffix.lower(),
            "file_size": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        }
    
    def _load_pdf(self, path: Path) -> List[Document]:
        """PDF dosyası yükle."""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(str(path))
            documents = []
            
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    metadata = self._get_base_metadata(path)
                    metadata["page_number"] = page_num
                    metadata["total_pages"] = len(reader.pages)
                    
                    documents.append(Document(content=text, metadata=metadata))
            
            return documents
            
        except ImportError:
            # Fallback to pdfplumber
            try:
                import pdfplumber
                
                documents = []
                with pdfplumber.open(str(path)) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        text = page.extract_text()
                        if text and text.strip():
                            metadata = self._get_base_metadata(path)
                            metadata["page_number"] = page_num
                            metadata["total_pages"] = len(pdf.pages)
                            
                            documents.append(Document(content=text, metadata=metadata))
                
                return documents
            except ImportError:
                raise ImportError("PDF yüklemek için pypdf veya pdfplumber gerekli")
    
    def _load_docx(self, path: Path) -> List[Document]:
        """Word dosyası yükle."""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(str(path))
            
            # Extract all paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_content.append(" | ".join(row_text))
                tables_text.append("\n".join(table_content))
            
            # Combine content
            content = "\n\n".join(paragraphs)
            if tables_text:
                content += "\n\n" + "\n\n".join(tables_text)
            
            metadata = self._get_base_metadata(path)
            
            return [Document(content=content, metadata=metadata)]
            
        except ImportError:
            raise ImportError("DOCX yüklemek için python-docx gerekli")
    
    def _load_excel(self, path: Path) -> List[Document]:
        """Excel dosyası yükle."""
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(str(path), data_only=True)
            documents = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(v.strip() for v in row_values):
                        rows.append(" | ".join(row_values))
                
                if rows:
                    content = "\n".join(rows)
                    metadata = self._get_base_metadata(path)
                    metadata["sheet_name"] = sheet_name
                    
                    documents.append(Document(content=content, metadata=metadata))
            
            return documents
            
        except ImportError:
            raise ImportError("Excel yüklemek için openpyxl gerekli")
    
    def _load_csv(self, path: Path) -> List[Document]:
        """CSV dosyası yükle."""
        import csv
        
        rows = []
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                if any(cell.strip() for cell in row):
                    rows.append(" | ".join(row))
        
        content = "\n".join(rows)
        metadata = self._get_base_metadata(path)
        
        return [Document(content=content, metadata=metadata)]
    
    def _load_text(self, path: Path) -> List[Document]:
        """Text/Markdown dosyası yükle."""
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        
        metadata = self._get_base_metadata(path)
        
        return [Document(content=content, metadata=metadata)]
    
    def _load_html(self, path: Path) -> List[Document]:
        """HTML dosyası yükle."""
        try:
            from bs4 import BeautifulSoup
            
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            
            # Remove scripts and styles
            for script in soup(["script", "style"]):
                script.decompose()
            
            content = soup.get_text(separator="\n", strip=True)
            metadata = self._get_base_metadata(path)
            
            # Extract title if exists
            title_tag = soup.find("title")
            if title_tag:
                metadata["title"] = title_tag.get_text(strip=True)
            
            return [Document(content=content, metadata=metadata)]
            
        except ImportError:
            raise ImportError("HTML yüklemek için beautifulsoup4 gerekli")
    
    def _load_json(self, path: Path) -> List[Document]:
        """JSON dosyası yükle."""
        import json
        
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        # Convert to readable text
        content = json.dumps(data, indent=2, ensure_ascii=False)
        metadata = self._get_base_metadata(path)
        
        return [Document(content=content, metadata=metadata)]


# Singleton instance
document_loader = DocumentLoader()
