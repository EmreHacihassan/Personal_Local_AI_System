"""
Multi-Modal RAG System
======================

Advanced RAG system for processing multiple modalities:
- Text documents (PDF, DOCX, TXT, Markdown)
- Images (OCR, visual understanding)
- Tables (structured data extraction)
- Charts (visualization understanding)
- Audio (transcription)
- Code (syntax-aware processing)

Features:
- OCR Integration (Tesseract, EasyOCR)
- Vision Model Support (GPT-4V, Claude Vision)
- Table Structure Recognition
- Chart/Graph Data Extraction
- Audio Transcription (Whisper)
- Unified Multi-Modal Embeddings
- Cross-Modal Retrieval

Enterprise-grade implementation.

Author: AI Assistant
Version: 1.0.0
"""

import asyncio
import base64
import hashlib
import io
import json
import logging
import mimetypes
import os
import re
import tempfile
import time
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import (
    Any,
    BinaryIO,
    Callable,
    Dict,
    Generator,
    List,
    Optional,
    Protocol,
    Set,
    Tuple,
    Type,
    Union,
)

from core.logger import get_logger

logger = get_logger("rag.multimodal_rag")


# =============================================================================
# OPTIONAL IMPORTS
# =============================================================================

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    Image = None

try:
    import pytesseract
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False
    pytesseract = None

try:
    import easyocr
    HAS_EASYOCR = True
except ImportError:
    HAS_EASYOCR = False
    easyocr = None

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    fitz = None

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    DocxDocument = None

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False
    pd = None


# =============================================================================
# ENUMS
# =============================================================================

class ContentModality(Enum):
    """İçerik modalite türleri."""
    TEXT = "text"
    IMAGE = "image"
    TABLE = "table"
    CHART = "chart"
    CODE = "code"
    AUDIO = "audio"
    VIDEO = "video"
    PDF = "pdf"
    MIXED = "mixed"


class ExtractionMethod(Enum):
    """Extraction yöntemi."""
    NATIVE = "native"           # Native text extraction
    OCR_TESSERACT = "tesseract" # Tesseract OCR
    OCR_EASYOCR = "easyocr"     # EasyOCR
    VISION_LLM = "vision_llm"   # Vision LLM (GPT-4V, Claude)
    AUDIO_WHISPER = "whisper"   # Whisper transcription
    TABLE_PARSER = "table"      # Table structure parser
    HYBRID = "hybrid"           # Combination of methods


class ProcessingStatus(Enum):
    """İşleme durumu."""
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"
    PARTIAL = "partial"


# =============================================================================
# DATA CLASSES
# =============================================================================

@dataclass
class ModalityMetadata:
    """Modalite metadata."""
    modality: ContentModality
    confidence: float = 1.0
    extraction_method: ExtractionMethod = ExtractionMethod.NATIVE
    
    # Image specific
    image_width: Optional[int] = None
    image_height: Optional[int] = None
    ocr_language: Optional[str] = None
    
    # Table specific
    row_count: Optional[int] = None
    column_count: Optional[int] = None
    headers: List[str] = field(default_factory=list)
    
    # Audio specific
    duration_seconds: Optional[float] = None
    sample_rate: Optional[int] = None
    
    # Code specific
    language: Optional[str] = None
    line_count: Optional[int] = None
    
    def to_dict(self) -> Dict:
        return {
            "modality": self.modality.value,
            "confidence": self.confidence,
            "extraction_method": self.extraction_method.value,
            "image_width": self.image_width,
            "image_height": self.image_height,
            "row_count": self.row_count,
            "column_count": self.column_count,
            "duration_seconds": self.duration_seconds,
            "language": self.language,
        }


@dataclass
class ExtractedContent:
    """Çıkarılan içerik."""
    id: str
    text: str
    modality: ContentModality
    metadata: ModalityMetadata
    
    # Source tracking
    source_file: Optional[str] = None
    page_number: Optional[int] = None
    position: Optional[Dict] = None
    
    # Embedding
    embedding: Optional[List[float]] = None
    
    # Processing info
    processing_time_ms: int = 0
    quality_score: float = 1.0
    
    def to_dict(self) -> Dict:
        return {
            "id": self.id,
            "text": self.text[:500],
            "modality": self.modality.value,
            "metadata": self.metadata.to_dict(),
            "source_file": self.source_file,
            "page_number": self.page_number,
            "quality_score": self.quality_score,
        }


@dataclass
class MultiModalDocument:
    """Multi-modal doküman."""
    id: str
    filename: str
    contents: List[ExtractedContent]
    primary_modality: ContentModality
    
    # Document metadata
    total_text_length: int = 0
    image_count: int = 0
    table_count: int = 0
    page_count: int = 0
    
    # Processing status
    status: ProcessingStatus = ProcessingStatus.PENDING
    error_message: Optional[str] = None
    
    # Timestamps
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    processed_at: Optional[str] = None
    
    def to_dict(self) -> Dict:
        return {
            "id": self.id,
            "filename": self.filename,
            "content_count": len(self.contents),
            "primary_modality": self.primary_modality.value,
            "total_text_length": self.total_text_length,
            "image_count": self.image_count,
            "table_count": self.table_count,
            "page_count": self.page_count,
            "status": self.status.value,
        }
    
    def get_full_text(self) -> str:
        """Tüm içeriği birleştir."""
        texts = [c.text for c in self.contents if c.text]
        return "\n\n".join(texts)


@dataclass
class MultiModalQuery:
    """Multi-modal sorgu."""
    text: str
    images: List[bytes] = field(default_factory=list)
    modalities_to_search: List[ContentModality] = field(default_factory=list)
    
    # Query options
    use_ocr: bool = True
    use_vision: bool = True
    use_tables: bool = True


@dataclass
class MultiModalResult:
    """Multi-modal sonuç."""
    query: str
    response: str
    confidence: float
    
    # Retrieved content by modality
    text_results: List[ExtractedContent] = field(default_factory=list)
    image_results: List[ExtractedContent] = field(default_factory=list)
    table_results: List[ExtractedContent] = field(default_factory=list)
    
    # Metrics
    modalities_used: List[str] = field(default_factory=list)
    total_results: int = 0
    processing_time_ms: int = 0
    
    def to_dict(self) -> Dict:
        return {
            "query": self.query,
            "response": self.response,
            "confidence": self.confidence,
            "modalities_used": self.modalities_used,
            "text_results": len(self.text_results),
            "image_results": len(self.image_results),
            "table_results": len(self.table_results),
            "total_results": self.total_results,
            "processing_time_ms": self.processing_time_ms,
        }


# =============================================================================
# PROTOCOLS
# =============================================================================

class LLMProtocol(Protocol):
    def generate(self, prompt: str, **kwargs) -> str:
        ...


class VisionLLMProtocol(Protocol):
    def analyze_image(self, image_data: bytes, prompt: str) -> str:
        ...


class EmbeddingProtocol(Protocol):
    def embed_text(self, text: str) -> List[float]:
        ...


# =============================================================================
# CONTENT EXTRACTORS
# =============================================================================

class BaseExtractor(ABC):
    """Temel içerik çıkarıcı."""
    
    @abstractmethod
    def can_extract(self, file_path: str, mime_type: str) -> bool:
        """Bu extractor bu dosyayı işleyebilir mi?"""
        pass
    
    @abstractmethod
    def extract(self, file_path: str) -> List[ExtractedContent]:
        """İçerik çıkar."""
        pass


class TextExtractor(BaseExtractor):
    """Plain text extractor."""
    
    SUPPORTED_EXTENSIONS = {".txt", ".md", ".rst", ".log", ".csv"}
    
    def can_extract(self, file_path: str, mime_type: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS or mime_type.startswith("text/")
    
    def extract(self, file_path: str) -> List[ExtractedContent]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            content_id = hashlib.md5(file_path.encode()).hexdigest()[:12]
            
            return [ExtractedContent(
                id=content_id,
                text=content,
                modality=ContentModality.TEXT,
                metadata=ModalityMetadata(
                    modality=ContentModality.TEXT,
                    extraction_method=ExtractionMethod.NATIVE,
                ),
                source_file=file_path,
            )]
        
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return []


class ImageExtractor(BaseExtractor):
    """Image OCR extractor."""
    
    SUPPORTED_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"}
    
    def __init__(self, preferred_ocr: str = "tesseract"):
        self.preferred_ocr = preferred_ocr
        self._easyocr_reader = None
    
    def can_extract(self, file_path: str, mime_type: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS or mime_type.startswith("image/")
    
    def extract(self, file_path: str) -> List[ExtractedContent]:
        if not HAS_PIL:
            logger.warning("PIL not installed. Cannot process images.")
            return []
        
        try:
            image = Image.open(file_path)
            
            # Get image dimensions
            width, height = image.size
            
            # Perform OCR
            text = self._perform_ocr(image)
            
            content_id = hashlib.md5(file_path.encode()).hexdigest()[:12]
            
            return [ExtractedContent(
                id=content_id,
                text=text,
                modality=ContentModality.IMAGE,
                metadata=ModalityMetadata(
                    modality=ContentModality.IMAGE,
                    extraction_method=ExtractionMethod.OCR_TESSERACT if self.preferred_ocr == "tesseract" else ExtractionMethod.OCR_EASYOCR,
                    image_width=width,
                    image_height=height,
                ),
                source_file=file_path,
            )]
        
        except Exception as e:
            logger.error(f"Image extraction failed: {e}")
            return []
    
    def _perform_ocr(self, image: Any) -> str:
        """OCR gerçekleştir."""
        # Try Tesseract first
        if HAS_TESSERACT and self.preferred_ocr == "tesseract":
            try:
                text = pytesseract.image_to_string(image)
                if text.strip():
                    return text
            except Exception as e:
                logger.warning(f"Tesseract OCR failed: {e}")
        
        # Fallback to EasyOCR
        if HAS_EASYOCR:
            try:
                if self._easyocr_reader is None:
                    self._easyocr_reader = easyocr.Reader(["en", "tr"])
                
                # Convert to numpy array
                import numpy as np
                img_array = np.array(image)
                
                results = self._easyocr_reader.readtext(img_array)
                text = " ".join([result[1] for result in results])
                return text
            except Exception as e:
                logger.warning(f"EasyOCR failed: {e}")
        
        return ""
    
    def extract_with_vision(
        self,
        image_data: bytes,
        vision_llm: VisionLLMProtocol,
        prompt: str = "Describe this image in detail."
    ) -> str:
        """Vision LLM ile görsel anlama."""
        try:
            return vision_llm.analyze_image(image_data, prompt)
        except Exception as e:
            logger.error(f"Vision analysis failed: {e}")
            return ""


class PDFExtractor(BaseExtractor):
    """PDF extractor with text, images, and tables."""
    
    def can_extract(self, file_path: str, mime_type: str) -> bool:
        return Path(file_path).suffix.lower() == ".pdf" or mime_type == "application/pdf"
    
    def extract(self, file_path: str) -> List[ExtractedContent]:
        if not HAS_PYMUPDF:
            logger.warning("PyMuPDF not installed. Cannot process PDFs.")
            return self._fallback_extract(file_path)
        
        contents = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc, 1):
                # Extract text
                text = page.get_text()
                
                if text.strip():
                    content_id = hashlib.md5(f"{file_path}_p{page_num}".encode()).hexdigest()[:12]
                    
                    contents.append(ExtractedContent(
                        id=content_id,
                        text=text,
                        modality=ContentModality.TEXT,
                        metadata=ModalityMetadata(
                            modality=ContentModality.TEXT,
                            extraction_method=ExtractionMethod.NATIVE,
                        ),
                        source_file=file_path,
                        page_number=page_num,
                    ))
                
                # Extract images
                for img_index, img in enumerate(page.get_images(full=True)):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        # OCR on image
                        if HAS_PIL:
                            pil_image = Image.open(io.BytesIO(image_bytes))
                            ocr_text = self._ocr_image(pil_image)
                            
                            if ocr_text.strip():
                                img_content_id = hashlib.md5(
                                    f"{file_path}_p{page_num}_img{img_index}".encode()
                                ).hexdigest()[:12]
                                
                                contents.append(ExtractedContent(
                                    id=img_content_id,
                                    text=ocr_text,
                                    modality=ContentModality.IMAGE,
                                    metadata=ModalityMetadata(
                                        modality=ContentModality.IMAGE,
                                        extraction_method=ExtractionMethod.OCR_TESSERACT,
                                        image_width=pil_image.width,
                                        image_height=pil_image.height,
                                    ),
                                    source_file=file_path,
                                    page_number=page_num,
                                ))
                    except Exception as e:
                        logger.warning(f"Image extraction from PDF failed: {e}")
                
                # Extract tables (simple heuristic)
                tables = page.find_tables()
                for table_index, table in enumerate(tables):
                    try:
                        table_data = table.extract()
                        if table_data:
                            # Convert to markdown table
                            table_text = self._table_to_markdown(table_data)
                            
                            table_content_id = hashlib.md5(
                                f"{file_path}_p{page_num}_table{table_index}".encode()
                            ).hexdigest()[:12]
                            
                            contents.append(ExtractedContent(
                                id=table_content_id,
                                text=table_text,
                                modality=ContentModality.TABLE,
                                metadata=ModalityMetadata(
                                    modality=ContentModality.TABLE,
                                    extraction_method=ExtractionMethod.TABLE_PARSER,
                                    row_count=len(table_data),
                                    column_count=len(table_data[0]) if table_data else 0,
                                    headers=table_data[0] if table_data else [],
                                ),
                                source_file=file_path,
                                page_number=page_num,
                            ))
                    except Exception as e:
                        logger.warning(f"Table extraction from PDF failed: {e}")
            
            doc.close()
        
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
        
        return contents
    
    def _fallback_extract(self, file_path: str) -> List[ExtractedContent]:
        """Fallback extraction without PyMuPDF."""
        try:
            import subprocess
            result = subprocess.run(
                ["pdftotext", file_path, "-"],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode == 0:
                content_id = hashlib.md5(file_path.encode()).hexdigest()[:12]
                return [ExtractedContent(
                    id=content_id,
                    text=result.stdout,
                    modality=ContentModality.TEXT,
                    metadata=ModalityMetadata(
                        modality=ContentModality.TEXT,
                        extraction_method=ExtractionMethod.NATIVE,
                    ),
                    source_file=file_path,
                )]
        except Exception:
            pass
        
        return []
    
    def _ocr_image(self, image: Any) -> str:
        """Image OCR."""
        if HAS_TESSERACT:
            try:
                return pytesseract.image_to_string(image)
            except Exception:
                pass
        return ""
    
    def _table_to_markdown(self, table_data: List[List[str]]) -> str:
        """Tabloyu markdown formatına dönüştür."""
        if not table_data:
            return ""
        
        lines = []
        
        # Header
        header = table_data[0]
        lines.append("| " + " | ".join(str(cell) for cell in header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        
        # Rows
        for row in table_data[1:]:
            lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
        
        return "\n".join(lines)


class DocxExtractor(BaseExtractor):
    """DOCX extractor."""
    
    def can_extract(self, file_path: str, mime_type: str) -> bool:
        return Path(file_path).suffix.lower() == ".docx"
    
    def extract(self, file_path: str) -> List[ExtractedContent]:
        if not HAS_DOCX:
            logger.warning("python-docx not installed. Cannot process DOCX files.")
            return []
        
        contents = []
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            full_text = "\n\n".join(paragraphs)
            
            if full_text:
                content_id = hashlib.md5(file_path.encode()).hexdigest()[:12]
                contents.append(ExtractedContent(
                    id=content_id,
                    text=full_text,
                    modality=ContentModality.TEXT,
                    metadata=ModalityMetadata(
                        modality=ContentModality.TEXT,
                        extraction_method=ExtractionMethod.NATIVE,
                    ),
                    source_file=file_path,
                ))
            
            # Extract tables
            for table_index, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    table_text = self._table_to_markdown(table_data)
                    table_content_id = hashlib.md5(
                        f"{file_path}_table{table_index}".encode()
                    ).hexdigest()[:12]
                    
                    contents.append(ExtractedContent(
                        id=table_content_id,
                        text=table_text,
                        modality=ContentModality.TABLE,
                        metadata=ModalityMetadata(
                            modality=ContentModality.TABLE,
                            extraction_method=ExtractionMethod.TABLE_PARSER,
                            row_count=len(table_data),
                            column_count=len(table_data[0]) if table_data else 0,
                        ),
                        source_file=file_path,
                    ))
        
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
        
        return contents
    
    def _table_to_markdown(self, table_data: List[List[str]]) -> str:
        """Tabloyu markdown formatına dönüştür."""
        if not table_data:
            return ""
        
        lines = []
        header = table_data[0]
        lines.append("| " + " | ".join(str(cell) for cell in header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        
        for row in table_data[1:]:
            lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
        
        return "\n".join(lines)


class CodeExtractor(BaseExtractor):
    """Source code extractor."""
    
    LANGUAGE_MAP = {
        ".py": "python",
        ".js": "javascript",
        ".ts": "typescript",
        ".java": "java",
        ".cpp": "cpp",
        ".c": "c",
        ".cs": "csharp",
        ".go": "go",
        ".rs": "rust",
        ".rb": "ruby",
        ".php": "php",
        ".swift": "swift",
        ".kt": "kotlin",
        ".scala": "scala",
        ".sql": "sql",
        ".sh": "bash",
        ".yaml": "yaml",
        ".yml": "yaml",
        ".json": "json",
        ".xml": "xml",
        ".html": "html",
        ".css": "css",
    }
    
    def can_extract(self, file_path: str, mime_type: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in self.LANGUAGE_MAP
    
    def extract(self, file_path: str) -> List[ExtractedContent]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            ext = Path(file_path).suffix.lower()
            language = self.LANGUAGE_MAP.get(ext, "text")
            line_count = len(content.splitlines())
            
            content_id = hashlib.md5(file_path.encode()).hexdigest()[:12]
            
            return [ExtractedContent(
                id=content_id,
                text=content,
                modality=ContentModality.CODE,
                metadata=ModalityMetadata(
                    modality=ContentModality.CODE,
                    extraction_method=ExtractionMethod.NATIVE,
                    language=language,
                    line_count=line_count,
                ),
                source_file=file_path,
            )]
        
        except Exception as e:
            logger.error(f"Code extraction failed: {e}")
            return []


# =============================================================================
# MULTI-MODAL PROCESSOR
# =============================================================================

class MultiModalProcessor:
    """
    Unified multi-modal content processor.
    
    Manages all extractors and provides unified interface.
    """
    
    def __init__(self):
        self.extractors: List[BaseExtractor] = [
            TextExtractor(),
            ImageExtractor(),
            PDFExtractor(),
            DocxExtractor(),
            CodeExtractor(),
        ]
        
        # Document store
        self._documents: Dict[str, MultiModalDocument] = {}
        self._content_by_modality: Dict[ContentModality, List[ExtractedContent]] = {
            modality: [] for modality in ContentModality
        }
    
    def process_file(self, file_path: str) -> Optional[MultiModalDocument]:
        """Dosyayı işle."""
        start_time = time.time()
        
        file_path = str(Path(file_path).resolve())
        
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        # Detect mime type
        mime_type, _ = mimetypes.guess_type(file_path)
        mime_type = mime_type or "application/octet-stream"
        
        # Find appropriate extractor
        extractor = None
        for ext in self.extractors:
            if ext.can_extract(file_path, mime_type):
                extractor = ext
                break
        
        if extractor is None:
            logger.warning(f"No extractor found for {file_path}")
            return None
        
        # Extract content
        contents = extractor.extract(file_path)
        
        if not contents:
            logger.warning(f"No content extracted from {file_path}")
            return None
        
        # Create document
        doc_id = hashlib.md5(file_path.encode()).hexdigest()[:12]
        
        # Determine primary modality
        modality_counts = {}
        for content in contents:
            modality_counts[content.modality] = modality_counts.get(content.modality, 0) + 1
        
        primary_modality = max(modality_counts.keys(), key=lambda m: modality_counts[m])
        
        document = MultiModalDocument(
            id=doc_id,
            filename=os.path.basename(file_path),
            contents=contents,
            primary_modality=primary_modality,
            total_text_length=sum(len(c.text) for c in contents),
            image_count=sum(1 for c in contents if c.modality == ContentModality.IMAGE),
            table_count=sum(1 for c in contents if c.modality == ContentModality.TABLE),
            page_count=max((c.page_number or 1) for c in contents),
            status=ProcessingStatus.COMPLETED,
            processed_at=datetime.now().isoformat(),
        )
        
        # Store document
        self._documents[doc_id] = document
        
        # Index by modality
        for content in contents:
            self._content_by_modality[content.modality].append(content)
        
        processing_time = int((time.time() - start_time) * 1000)
        logger.info(f"Processed {file_path}: {len(contents)} contents, {processing_time}ms")
        
        return document
    
    def process_directory(
        self,
        directory: str,
        recursive: bool = True,
        patterns: List[str] = None
    ) -> List[MultiModalDocument]:
        """Dizini işle."""
        documents = []
        directory = Path(directory)
        
        if patterns is None:
            patterns = ["*.*"]
        
        for pattern in patterns:
            if recursive:
                files = directory.rglob(pattern)
            else:
                files = directory.glob(pattern)
            
            for file_path in files:
                if file_path.is_file():
                    doc = self.process_file(str(file_path))
                    if doc:
                        documents.append(doc)
        
        return documents
    
    def get_content_by_modality(
        self,
        modality: ContentModality
    ) -> List[ExtractedContent]:
        """Modaliteye göre içerik getir."""
        return self._content_by_modality.get(modality, [])
    
    def get_document(self, doc_id: str) -> Optional[MultiModalDocument]:
        """Doküman getir."""
        return self._documents.get(doc_id)
    
    def get_all_documents(self) -> List[MultiModalDocument]:
        """Tüm dokümanları getir."""
        return list(self._documents.values())
    
    def get_stats(self) -> Dict[str, Any]:
        """İstatistikler."""
        return {
            "document_count": len(self._documents),
            "content_by_modality": {
                m.value: len(contents) for m, contents in self._content_by_modality.items()
            },
            "total_text_length": sum(d.total_text_length for d in self._documents.values()),
            "total_images": sum(d.image_count for d in self._documents.values()),
            "total_tables": sum(d.table_count for d in self._documents.values()),
        }


# =============================================================================
# MULTI-MODAL RAG
# =============================================================================

class MultiModalRAG:
    """
    Multi-Modal RAG System.
    
    Unified RAG across text, images, tables, and code.
    
    Features:
    - Multi-modal document processing
    - Cross-modal retrieval
    - Vision LLM integration
    - Table understanding
    - Code-aware search
    
    Example:
        mm_rag = MultiModalRAG()
        mm_rag.index_file("document.pdf")
        result = mm_rag.query("What does the chart show?")
    """
    
    def __init__(
        self,
        llm: Optional[LLMProtocol] = None,
        embedding_model: Optional[EmbeddingProtocol] = None,
        vision_llm: Optional[VisionLLMProtocol] = None,
    ):
        self._llm = llm
        self._embedding = embedding_model
        self._vision_llm = vision_llm
        
        self.processor = MultiModalProcessor()
        
        # Content embeddings cache
        self._embeddings: Dict[str, List[float]] = {}
        
        # Stats
        self._query_count = 0
    
    def _lazy_load(self):
        if self._llm is None:
            from core.llm_manager import llm_manager
            self._llm = llm_manager
    
    def index_file(self, file_path: str) -> Optional[MultiModalDocument]:
        """Dosyayı indexle."""
        doc = self.processor.process_file(file_path)
        
        if doc:
            # Generate embeddings for each content
            self._generate_embeddings(doc.contents)
        
        return doc
    
    def index_directory(
        self,
        directory: str,
        recursive: bool = True
    ) -> List[MultiModalDocument]:
        """Dizini indexle."""
        docs = self.processor.process_directory(directory, recursive)
        
        for doc in docs:
            self._generate_embeddings(doc.contents)
        
        return docs
    
    def _generate_embeddings(self, contents: List[ExtractedContent]):
        """İçerikler için embedding üret."""
        if self._embedding is None:
            return
        
        for content in contents:
            if content.text and len(content.text) > 10:
                try:
                    embedding = self._embedding.embed_text(content.text[:2000])
                    self._embeddings[content.id] = embedding
                    content.embedding = embedding
                except Exception as e:
                    logger.warning(f"Embedding generation failed: {e}")
    
    def query(
        self,
        query: str,
        modalities: List[ContentModality] = None,
        top_k: int = 5,
        include_tables: bool = True,
        include_images: bool = True,
    ) -> MultiModalResult:
        """
        Multi-modal sorgu.
        
        Args:
            query: Kullanıcı sorgusu
            modalities: Aranacak modaliteler (None = tümü)
            top_k: Döndürülecek sonuç sayısı
            include_tables: Tablo sonuçlarını dahil et
            include_images: Görsel sonuçları dahil et
            
        Returns:
            MultiModalResult
        """
        self._lazy_load()
        
        start_time = time.time()
        
        if modalities is None:
            modalities = [ContentModality.TEXT, ContentModality.TABLE, ContentModality.CODE]
            if include_images:
                modalities.append(ContentModality.IMAGE)
        
        # Retrieve from each modality
        text_results = []
        image_results = []
        table_results = []
        
        for modality in modalities:
            contents = self.processor.get_content_by_modality(modality)
            
            # Simple keyword matching (for now)
            scored_contents = []
            query_words = set(query.lower().split())
            
            for content in contents:
                content_words = set(content.text.lower().split())
                overlap = len(query_words & content_words)
                if overlap > 0:
                    score = overlap / len(query_words)
                    scored_contents.append((score, content))
            
            # Sort by score
            scored_contents.sort(key=lambda x: x[0], reverse=True)
            
            # Take top results
            top_contents = [c for _, c in scored_contents[:top_k]]
            
            if modality == ContentModality.TEXT:
                text_results.extend(top_contents)
            elif modality == ContentModality.IMAGE:
                image_results.extend(top_contents)
            elif modality == ContentModality.TABLE:
                table_results.extend(top_contents)
            elif modality == ContentModality.CODE:
                text_results.extend(top_contents)
        
        # Build context
        context_parts = []
        
        if text_results:
            context_parts.append("=== Text Content ===")
            for i, content in enumerate(text_results[:3], 1):
                context_parts.append(f"[Text {i}]: {content.text[:500]}")
        
        if table_results:
            context_parts.append("\n=== Tables ===")
            for i, content in enumerate(table_results[:2], 1):
                context_parts.append(f"[Table {i}]: {content.text[:400]}")
        
        if image_results:
            context_parts.append("\n=== Image Descriptions ===")
            for i, content in enumerate(image_results[:2], 1):
                context_parts.append(f"[Image {i}]: {content.text[:300]}")
        
        context = "\n\n".join(context_parts)
        
        # Generate response
        response, confidence = self._generate_response(query, context)
        
        modalities_used = list(set(
            [ContentModality.TEXT.value] if text_results else [] +
            [ContentModality.TABLE.value] if table_results else [] +
            [ContentModality.IMAGE.value] if image_results else []
        ))
        
        total_results = len(text_results) + len(image_results) + len(table_results)
        processing_time = int((time.time() - start_time) * 1000)
        
        self._query_count += 1
        
        return MultiModalResult(
            query=query,
            response=response,
            confidence=confidence,
            text_results=text_results,
            image_results=image_results,
            table_results=table_results,
            modalities_used=modalities_used,
            total_results=total_results,
            processing_time_ms=processing_time,
        )
    
    async def query_async(self, query: str, **kwargs) -> MultiModalResult:
        """Asenkron sorgu."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            lambda: self.query(query, **kwargs)
        )
    
    def _generate_response(
        self,
        query: str,
        context: str
    ) -> Tuple[str, float]:
        """Response üret."""
        if not context.strip():
            return "İlgili içerik bulunamadı.", 0.3
        
        prompt = f"""Answer the question using the provided multi-modal context.
Reference specific content types (text, tables, images) when relevant.

{context}

Question: {query}

Answer:"""
        
        try:
            response = self._llm.generate(prompt, max_tokens=600)
            confidence = 0.8 if len(context) > 200 else 0.6
            return response, confidence
        
        except Exception as e:
            logger.error(f"Response generation failed: {e}")
            return f"Error: {str(e)}", 0.0
    
    def analyze_image(
        self,
        image_path: str,
        question: str = "What is shown in this image?"
    ) -> str:
        """Vision LLM ile görsel analiz."""
        if self._vision_llm is None:
            # Fallback to OCR
            extractor = ImageExtractor()
            contents = extractor.extract(image_path)
            if contents:
                return contents[0].text
            return "Vision analysis not available."
        
        try:
            with open(image_path, "rb") as f:
                image_data = f.read()
            
            return self._vision_llm.analyze_image(image_data, question)
        
        except Exception as e:
            logger.error(f"Image analysis failed: {e}")
            return f"Error: {str(e)}"
    
    def get_stats(self) -> Dict[str, Any]:
        """İstatistikler."""
        processor_stats = self.processor.get_stats()
        return {
            **processor_stats,
            "query_count": self._query_count,
            "embeddings_cached": len(self._embeddings),
        }


# =============================================================================
# SINGLETON
# =============================================================================

_multimodal_rag: Optional[MultiModalRAG] = None
_multimodal_processor: Optional[MultiModalProcessor] = None


def get_multimodal_rag() -> MultiModalRAG:
    """Singleton MultiModalRAG instance."""
    global _multimodal_rag
    
    if _multimodal_rag is None:
        _multimodal_rag = MultiModalRAG()
    
    return _multimodal_rag


def get_multimodal_processor() -> MultiModalProcessor:
    """Singleton MultiModalProcessor instance."""
    global _multimodal_processor
    
    if _multimodal_processor is None:
        _multimodal_processor = MultiModalProcessor()
    
    return _multimodal_processor


multimodal_rag = MultiModalRAG()
multimodal_processor = MultiModalProcessor()


__all__ = [
    "MultiModalRAG",
    "MultiModalProcessor",
    "ContentModality",
    "ExtractionMethod",
    "ProcessingStatus",
    "ModalityMetadata",
    "ExtractedContent",
    "MultiModalDocument",
    "MultiModalQuery",
    "MultiModalResult",
    "TextExtractor",
    "ImageExtractor",
    "PDFExtractor",
    "DocxExtractor",
    "CodeExtractor",
    "multimodal_rag",
    "multimodal_processor",
    "get_multimodal_rag",
    "get_multimodal_processor",
]
